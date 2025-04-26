import streamlit as st
from main import run_research  
from draft_agent import format_citation, STYLE_TEMPLATES  
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT
import io
import requests
import datetime
import logging
import re
from docx import Document
from docx.shared import Pt, Inches

# Set up logging
logging.basicConfig(filename="research_agent.log", level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

# Initialize all session state variables at the top
# Define all your session-state defaults in one place
DEFAULTS = {
    "research_data": None,
    "response": None,
    "pdf_buffer": None,
    "word_buffer": None,
    "writing_style": "Academic",
    "language": "English",
    "citation_format": "APA",
    "target_word_count": 1000,
}

# Initialize any missing keys in st.session_state
for key, default in DEFAULTS.items():
    if key not in st.session_state:
        st.session_state[key] = default

# Inject custom CSS for improved readability and aesthetics
st.markdown("""
    <style>
    /* Base styling - Futuristic Dark Theme */
    .stApp {
        background-color: #0d1117;
        color: #f2e6ea;
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
        line-height: 1.6;
    }
    
    /* Typography */
    h1 {
        color: #ff6f8a;  /* Soft light pink */
        font-size: 2.5rem;
        font-weight: 700;
        letter-spacing: -0.02em;
        margin-bottom: 1.5rem;
        border-bottom: 2px solid #e68ca0;  /* Slightly deeper pink for contrast */
        padding-bottom: 0.6rem;
        text-shadow: 0 0 15px rgba(255, 182, 193, 0.5);  /* Pink glow */
    }
    
    h2 {
        color: #ff6f8a;        /* Soft rose-pink that pairs with #ffb6c1 */
        font-size: 1.8rem;
        font-weight: 600;
        margin-top: 1.8rem;
        margin-bottom: 1rem;
    }

    
    h3 {
        color: #ff6f8a;       /* Soft pale rose-pink matching with theme */
        font-size: 1.4rem;
        font-weight: 500;
        margin-top: 1.5rem;
        margin-bottom: 0.8rem;
    }
    
    p, li {
        font-size: 1rem;
        color: #e0a9a9;  /* Soft pink-gray shade */
        line-height: 1.7;
    }
    
    /* Text input styling */
    .stTextInput > div > input {
        background-color: #161b22;
        color: #e0a9a9;                           /* Soft pink-gray text */
        border: 1px solid #362c30;                /* Muted rose-gray border */
        border-radius: 8px;
        box-shadow: 0 0 5px rgba(255, 182, 193, 0.2); /* Gentle pink glow */
        padding: 12px 16px;
        font-size: 1rem;
        transition: all 0.2s ease;
    }

    .stTextInput > div > input:focus {
        border-color: #ff9db1;                    /* Brighter pink on focus */
        box-shadow: 0 0 10px rgba(255, 182, 193, 0.4);
        outline: none;
    }

    /* Button styling (for primary action buttons) */
    .stButton > button {
        background: linear-gradient(90deg, #ff6b8a, #ff4f73) ;  /* Soft pink gradient */
        color: #0d1117;                                       /* Dark text for contrast */
        font-weight: 600;
        padding: 0.6rem 1.5rem;
        border: none;
        border-radius: 8px;
        cursor: pointer;
        transition: all 0.3s ease;
        box-shadow: 0 0 15px rgba(255, 182, 193, 0.4);        /* Pink glow */
        position: relative;
        overflow: hidden;
    }

    .stButton > button:hover {
        box-shadow: 0 0 20px rgba(255, 182, 193, 0.6);
        transform: translateY(-2px);
        background: linear-gradient(45deg, #ffb6c1, #ff9db1);  /* Reverse gradient on hover */
    }

    .stButton > button:active {
        transform: translateY(1px);
        box-shadow: 0 0 10px rgba(255, 182, 193, 0.3);
    }

    .stButton > button::after {
        content: '';
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: rgba(255, 255, 255, 0.1);
        transform: rotate(30deg);
        transition: transform 0.3s ease;
    }

    .stButton > button:hover::after {
        transform: rotate(30deg) translate(-10%, -10%);
    }

    
    /* Download button styling (for Downloading PDF... 📄 button) */
    .stDownloadButton > button {
        background: linear-gradient(90deg, #ff6b8a, #ff4f73);  /* Soft pink gradient */
        color: #0d1117;                                       /* Dark text for contrast */
        font-weight: 600;
        padding: 0.7rem 2rem;
        border-radius: 8px;
        border: none;
        cursor: pointer;
        transition: all 0.3s ease;
        box-shadow: 0 0 15px rgba(255, 182, 193, 0.5);        /* Pink glow */
        position: relative;
        overflow: hidden;
        display: inline-flex;
        align-items: center;
        gap: 0.5rem;
        font-size: 1rem;
        margin-top: 0.5rem;
    }

    .stDownloadButton > button:hover {
        box-shadow: 0 0 25px rgba(255, 182, 193, 0.7);
        transform: translateY(-2px) scale(1.05);
        background: linear-gradient(45deg, #ffb6c1, #ff9db1);  /* Reverse gradient on hover */
    }

    .stDownloadButton > button:active {
        transform: translateY(1px) scale(0.98);
        box-shadow: 0 0 10px rgba(255, 182, 193, 0.3);
    }

    .stDownloadButton > button::after {
        content: '';
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: rgba(255, 255, 255, 0.1);
        transform: rotate(30deg);
        transition: transform 0.3s ease;
    }

    .stDownloadButton > button:hover::after {
        transform: rotate(30deg) translate(-10%, -10%);
    }

    
    /* Custom Selectbox styling */
    .custom-select-wrapper {
        position: relative;
        width: 100%;
        font-size: 1rem;
    }

    .custom-select {
        background: linear-gradient(90deg, #ff6b8a, #ff4f73);  /* Soft pink gradient */
        color: #0d1117;                                       /* Dark text for contrast */
        font-weight: 600;
        border: none;
        border-radius: 8px;
        padding: 10px;
        font-size: 1rem;
        transition: all 0.3s ease;
        box-shadow: 0 0 10px rgba(255, 182, 193, 0.3);        /* Pink glow */
        width: 100%;
        cursor: pointer;
        appearance: none;
        -webkit-appearance: none;
        -moz-appearance: none;
    }

    .custom-select:hover {
        box-shadow: 0 0 15px rgba(255, 182, 193, 0.5);
        transform: translateY(-2px) scale(1.02);
        background: linear-gradient(45deg, #ffb6c1, #ff9db1);  /* Reverse gradient on hover */
    }

    .custom-select:focus {
        outline: none;
        box-shadow: 0 0 15px rgba(255, 182, 193, 0.5);
    }

    /* Custom arrow for the dropdown */
    .custom-select-wrapper::after {
        content: '\\25BC'; /* Unicode for down arrow */
        position: absolute;
        right: 15px;
        top: 50%;
        transform: translateY(-50%);
        color: #0d1117;  /* Dark arrow for better contrast */
        font-size: 1rem;
        pointer-events: none;
    }

    
    /* Dropdown menu styling */
    .custom-select option {
        background: #ffebf0;  /* Soft light pink background */
        color: #0d1117;       /* Dark text for contrast */
        font-weight: 500;
        border: 1px solid #ff9db1;  /* Light pink border */
        border-radius: 8px;
        box-shadow: 0 0 5px rgba(255, 182, 193, 0.2);  /* Soft pink glow */
        padding: 10px;
    }

    .custom-select option:hover {
        background: linear-gradient(45deg, #ff9db1, #ffb6c1);  /* Soft pink gradient on hover */
        color: #0d1117;  /* Dark text on hover */
        box-shadow: 0 0 10px rgba(255, 182, 193, 0.3);  /* Soft pink glow on hover */
    }

    
    /* Hide the default st.selectbox */
    .stSelectbox {
        display: none !important;
    }

    /* Sidebar styling */
    .stSidebar {
        background-color: #f4f4f9;  /* Soft light background for the sidebar */
        border-right: 1px solid #e0e0e0;  /* Light border for a softer look */
    }

    .stSidebar .stMarkdown {
        color: #2d2d2d !important;  /* Dark text for readability */
    }

    .stSidebar h2 {
        color: #f06292;  /* Soft pink color for the sidebar header */
        font-size: 1.4rem;
        margin-top: 2rem;
        margin-bottom: 1rem;
        padding-bottom: 0.5rem;
        border-bottom: 1px solid #e0e0e0;  /* Light border to match sidebar */
        text-shadow: 0 0 10px rgba(240, 98, 146, 0.3);  /* Soft pink glow effect */
    }

    /* Status indicators in sidebar */
    .stSidebar .stSuccess {
        background: linear-gradient(90deg, #ff6b8a, #ff4f73) !important; /* Soft pink gradient */
        color: #0d1117 !important;                                        /* Dark text for contrast */
        padding: 10px 15px;
        border-radius: 8px;
        font-weight: 500;
        margin-bottom: 1rem;
        display: flex;
        align-items: center;
        box-shadow: 0 0 10px rgba(255, 182, 193, 0.4);                     /* Pink glow */
        border-left: 3px solid #ff4081;                                    /* Vibrant pink accent */
    }

    .stSidebar .stError {
        background: linear-gradient(90deg, #ff8a80, #ff5252) !important;  /* Warm coral-red gradient */
        color: #0d1117 !important;                                        /* Dark text for contrast */
        padding: 10px 15px;
        border-radius: 8px;
        font-weight: 500;
        margin-bottom: 1rem;
        display: flex;
        align-items: center;
        box-shadow: 0 0 10px rgba(255, 82, 82, 0.4);                       /* Soft red glow */
        border-left: 3px solid #e91e63;                                   /* Deep pink-red accent */
    }

    
    /* JSON/Code block styling */
    .stCodeBlock {
        background-color: #1e1e2f;  /* Darker background for code blocks */
        color: #f1f1f1;              /* Light color text for better readability */
        border-radius: 8px;
        padding: 1.2rem;
        font-family: 'Fira Code', 'JetBrains Mono', 'Consolas', monospace;
        font-size: 0.9rem;
        line-height: 1.6;
        overflow-x: auto;
        border: 1px solid #44475a;    /* Darker border for subtle contrast */
        box-shadow: inset 0 0 10px rgba(0, 0, 0, 0.3);  /* More prominent shadow for depth */
    }

    /* Progress bar */
    .stProgress > div > div {
        background-color: #21262d;   /* Dark background for the progress bar */
        height: 0.6rem !important;
        border-radius: 1rem;
    }

    .stProgress > div > div > div {
        background: linear-gradient(90deg, #ff6b8a, #ff4f73) !important; /* Soft pink gradient for progress */
        border-radius: 1rem;
        box-shadow: 0 0 10px rgba(255, 128, 128, 0.5);  /* Soft glow around the progress */
    }

    /* Notification messages */
    .stSuccess {
        background-color: rgba(89, 154, 156, 0.2) !important; /* Light greenish background */
        color: #34d399 !important;  /* Soft green color */
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #10b981; /* Green border */
        margin: 1rem 0;
        font-size: 1rem;
        box-shadow: 0 0 10px rgba(6, 95, 70, 0.2); /* Subtle shadow */
    }

    .stInfo {
        background-color: rgba(255, 182, 193, 0.2) !important; /* Soft pink background */
        color: #ff99cc !important;  /* Light pink color */
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #ff66b2; /* Pink border */
        margin: 1rem 0;
        font-size: 1rem;
        box-shadow: 0 0 10px rgba(255, 182, 193, 0.2); /* Soft pink shadow */
    }

    .stWarning {
        background-color: rgba(255, 160, 122, 0.2) !important; /* Light peach background */
        color: #fbbf24 !important; /* Bright yellow color */
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #f59e0b; /* Yellow border */
        margin: 1rem 0;
        font-size: 1rem;
        box-shadow: 0 0 10px rgba(255, 160, 122, 0.2); /* Soft orange shadow */
    }

    
    /* Error message styling */
    .stError {
        background-color: rgba(255, 99, 71, 0.15) !important; /* Lighter soft red */
        color: #ff6b6b !important; /* Softer error text */
        padding: 1rem;
        border-radius: 10px;
        border-left: 4px solid #f87171; /* Bright red border */
        margin: 1.2rem 0;
        font-size: 1rem;
        box-shadow: 0 2px 12px rgba(255, 99, 71, 0.3); /* Slightly bigger shadow */
    }

    /* JSON display styling */
    .stJson {
        background-color: #0d1117; /* Darker background */
        border: 1px solid #2d333b; /* Softer border */
        border-radius: 10px;
        padding: 1.5rem;
        margin: 1.2rem 0;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.25); /* Deeper shadow */
        font-family: 'Fira Code', 'JetBrains Mono', 'Consolas', monospace;
        font-size: 0.95rem;
    }

    /* Download button container styling */
    .download-btn-container {
        margin-top: 1.5rem;
        text-align: center;
    }

    /* Format description styling */
    .format-description {
        margin-top: 0.7rem;
        font-size: 0.95rem;
        color: #7dcfff; /* Softer blue */
        text-align: center;
        transition: opacity 0.3s ease, transform 0.3s ease;
    }

    .format-description:hover {
        opacity: 0.85;
        transform: translateY(-2px);
    }

    
    /* Spinner */
    .stSpinner > div {
        border-color: #38bdf8 !important; /* Bright sky-blue */
        border-bottom-color: transparent !important;
        filter: drop-shadow(0 0 10px rgba(56, 189, 248, 0.7));
        animation: spin 1s linear infinite;
    }

    @keyframes spin {
        from { transform: rotate(0deg); }
        to { transform: rotate(360deg); }
    }

    /* Feedback textarea */
    .stTextArea > div > textarea {
        background-color: #0d1117; /* Slightly deeper dark */
        color: #d1d5db; /* Soft light gray */
        border: 1px solid #2d333b;
        border-radius: 10px;
        padding: 14px;
        font-size: 0.95rem;
        line-height: 1.6;
        resize: vertical;
        box-shadow: 0 0 6px rgba(59, 130, 246, 0.25);
        transition: border-color 0.3s ease, box-shadow 0.3s ease, background-color 0.3s ease;
    }

    .stTextArea > div > textarea:focus {
        background-color: #111827;
        border-color: #3b82f6;
        box-shadow: 0 0 12px rgba(59, 130, 246, 0.5);
        outline: none;
    }

    /* General spacing and container improvements */
    .main .block-container {
        padding: 2.5rem 2rem;
        max-width: 1000px;
        margin: 0 auto;
        background-color: rgba(255, 255, 255, 0.02);
        border-radius: 12px;
        box-shadow: 0 6px 20px rgba(0, 0, 0, 0.2);
    }

    
    /* Markdown block spacing */
    .stMarkdown {
        margin-bottom: 2rem;
        font-size: 1rem;
        line-height: 1.6;
        color: #c9d1d9;
    }

    /* Futuristic background glow */
    .stApp::before {
        content: "";
        position: fixed;
        top: 0;
        left: 0;
        width: 100%;
        height: 100%;
        background: radial-gradient(circle at top right, rgba(14, 165, 233, 0.07), transparent 70%);
        pointer-events: none;
        z-index: -1;
       animation: softPulse 8s ease-in-out infinite;
    }

    @keyframes softPulse {
        0%, 100% {
            opacity: 0.6;
            }
        50% {
            opacity: 0.9;
            }
    }

    /* Title underline animation */
    h1 {
        position: relative;
        padding-bottom: 6px;
        margin-bottom: 1.2rem;
        color: #e0f2fe;
    }

    h1::after {
        content: '';
        position: absolute;
        height: 3px;
        width: 0;
        left: 0;
        bottom: 0;
        background: linear-gradient(90deg, #0ea5e9, #38bdf8, transparent);
        border-radius: 2px;
        transition: width 0.5s ease-in-out;
    }

    h1:hover::after {
        width: 80%;
    }

    
    /* Enhanced Custom Scrollbar */
    ::-webkit-scrollbar {
        width: 10px;
        height: 10px;
    }

    ::-webkit-scrollbar-track {
        background: #0d1117;
        border-radius: 10px;
    }

    ::-webkit-scrollbar-thumb {
        background: linear-gradient(180deg, #2f3542, #3a3f47);
        border-radius: 10px;
        box-shadow: inset 0 0 4px rgba(88, 166, 255, 0.2);
        transition: background 0.3s ease, box-shadow 0.3s ease;
    }

    ::-webkit-scrollbar-thumb:hover {
        background: #58a6ff;
        box-shadow: 0 0 8px rgba(88, 166, 255, 0.5);
    }
    </style>
    """, unsafe_allow_html=True)


# Function to check OpenRouter status
def check_openrouter_status():
    """Check if OpenRouter API is operational."""
    try:
        response = requests.get("https://openrouter.ai/api/v1/models", timeout=5)
        return response.status_code == 200
    except Exception as e:
        logging.error(f"Failed to check OpenRouter status: {str(e)}")
        return False

# Function to add page numbers to the PDF
def on_page(canvas, doc):
    page_num = canvas.getPageNumber()
    text = f"Page {page_num}"
    canvas.saveState()
    canvas.setFont('Helvetica', 8)
    canvas.setFillColor(colors.grey)
    canvas.drawRightString(doc.rightMargin + doc.width, doc.bottomMargin - 10, text)
    canvas.restoreState()

def format_reference_for_pdf(ref_text):
    """Format a single reference for PDF."""
    # Extract date and URL using regex
    date_match = re.search(r'\(([0-9]{4},\s*[^)]+)\)', ref_text)
    url_match = re.search(r'Retrieved from\s+(https?://\S+)', ref_text)
    
    if date_match and url_match:
        date = date_match.group(1)
        url = url_match.group(1)
        return f"({date}). Retrieved from {url}"
    return ref_text

def preprocess_references(refs_section):
    """Preprocess references to ensure proper formatting."""
    # Split references by looking for date pattern and "Retrieved from"
    pattern = r'(?=\([0-9]{4},.*?\).*?Retrieved from)'
    refs = re.split(pattern, refs_section)
    # Clean up each reference
    refs = [ref.strip() for ref in refs if ref.strip()]
    
    formatted_refs = []
    for i, ref in enumerate(refs, 1):
        # Split URL if it's broken across lines
        ref = re.sub(r'(?<=https?://\S+)-\s*\n\s*(?=\S+)', '', ref)
        # Remove any "Page X" artifacts
        ref = re.sub(r'Page \d+\s*', '', ref)
        # Clean up extra whitespace
        ref = ' '.join(ref.split())
        formatted_ref = format_reference_for_pdf(ref, i)
        formatted_refs.append(formatted_ref)
    
    return formatted_refs

def format_references_section(refs_text):
    """Format references by adding line breaks after URLs."""
    # Split by "Retrieved from" to separate different references
    refs = refs_text.split("Retrieved from")
    formatted_refs = []
    
    for i, ref in enumerate(refs):
        if i == 0:  # First part might be empty
            continue
            
        # Find the URL pattern
        url_match = re.search(r'(https?://\S+)', ref)
        if url_match:
            url = url_match.group(1)
            # Get the text before the URL (title part)
            title_part = ref[:url_match.start()].strip()
            if i > 0 and formatted_refs:  # Add the previous title
                formatted_refs[-1] += f"Retrieved from {url}\n\n"
            formatted_refs.append(title_part)
    
    # Handle the last URL
    if formatted_refs and url_match:
        formatted_refs[-1] += f"Retrieved from {url}\n\n"
    
    return "".join(formatted_refs)

# Function to generate PDF with proper formatting and cover page
def generate_pdf(query, data, summary, deep_research=False):
    """Generate a PDF report with query, data, and summary in a research paper format."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        topMargin=72,
        bottomMargin=72,
        leftMargin=72,
        rightMargin=72
    )
    styles = getSampleStyleSheet()

    # Customize styles for a research paper look
    styles['Title'].fontSize = 16
    styles['Title'].spaceAfter = 12
    styles['Heading2'].fontSize = 14
    styles['Heading2'].spaceAfter = 6
    styles['Heading3'].fontSize = 12
    styles['Heading3'].spaceAfter = 6
    styles['BodyText'].fontSize = 10
    styles['BodyText'].leading = 14
    styles['BodyText'].spaceAfter = 12

    # Add a custom style for references
    styles.add(ParagraphStyle(
        name='Reference',
        parent=styles['BodyText'],
        fontSize=10,
        leftIndent=36,
        firstLineIndent=-36,
        spaceAfter=12
    ))

    story = []

    # Cover Page
    story.append(Paragraph("Deep Research AI Agent Report", styles['Title']))
    story.append(Spacer(1, 24))
    story.append(Paragraph(f"Query: {query}", styles['Normal']))
    story.append(Paragraph(f"Date: {datetime.date.today().strftime('%B %d, %Y')}", styles['Normal']))
    story.append(Paragraph(f"Author: [Your Name]", styles['Normal']))
    story.append(Spacer(1, 48))

    # Metadata
    story.append(Paragraph(f"OpenRouter Status: {'Operational' if check_openrouter_status() else 'Down'}", styles['Normal']))
    story.append(Spacer(1, 12))
    mode = "Deep Research" if deep_research else "Quick Research"
    story.append(Paragraph(f"Mode: {mode}", styles['Normal']))
    story.append(Spacer(1, 12))

    # Research Summary Section
    story.append(Paragraph("Research Summary", styles['Heading2']))
    story.append(Spacer(1, 12))

    # Add research data summary
    for item in data:
        story.append(Paragraph(f"• {item['title']}", styles['Heading3']))
        story.append(Paragraph(item['content'], styles['BodyText']))
        story.append(Spacer(1, 12))

    # Process the content sections
    sections = summary.split("\n\n")
    references = []
    current_heading = None
    
    for section in sections:
        if section.startswith("**") and section.endswith("**"):
            current_heading = section.strip("**").rstrip(":")
            if current_heading == "References":
                continue
            story.append(Paragraph(current_heading, styles['Heading2']))
            story.append(Spacer(1, 12))
        else:
            if current_heading == "References":
                # Collect references for later processing
                refs = section.split("\n")
                for ref in refs:
                    if ref.strip():
                        references.append(ref.strip())
            elif current_heading:
                if current_heading == "Analysis":
                    # Split into smaller paragraphs for readability
                    sentences = re.split(r'(?<=[.!?])\s+', section)
                    new_paragraphs = []
                    current_paragraph = []
                    sentence_count = 0
                    
                    max_sentences = 4 if len(sentences) < 15 else 6
                    
                    for sentence in sentences:
                        current_paragraph.append(sentence)
                        sentence_count += 1
                        if sentence_count >= max_sentences:
                            new_paragraphs.append(' '.join(current_paragraph))
                            current_paragraph = []
                            sentence_count = 0
                    
                    if current_paragraph:
                        new_paragraphs.append(' '.join(current_paragraph))
                    
                    for paragraph in new_paragraphs:
                        if paragraph.strip():
                            story.append(Paragraph(paragraph.strip(), styles['BodyText']))
                            story.append(Spacer(1, 12))
                elif current_heading == "Key Findings":
                    findings = re.split(r'(?=\d+\.)', section)
                    for finding in findings:
                        if finding.strip():
                            story.append(Paragraph(finding.strip(), styles['BodyText']))
                            story.append(Spacer(1, 6))
                else:
                    formatted_text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", section)
                    formatted_text = formatted_text.replace("**", "")
                    story.append(Paragraph(formatted_text, styles['BodyText']))
                    story.append(Spacer(1, 12))

    # Add References section at the end
    if references:
        story.append(Paragraph("References", styles['Heading2']))
        story.append(Spacer(1, 12))
        
        for i, ref in enumerate(references, 1):
            formatted_ref = format_reference_for_pdf(ref)
            if formatted_ref:
                ref_para = Paragraph(
                    f"{i}. {formatted_ref}",
                    styles['Reference']
                )
                story.append(ref_para)

    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)
    buffer.seek(0)
    return buffer

# Function to generate Word document
def generate_docx(query, data, summary, deep_research=False):
    """Generate a Word document with query, data, and summary."""
    doc = Document()
    doc.add_heading("Deep Research AI Agent Report", 0)
    doc.add_paragraph(f"Query: {query}")
    doc.add_paragraph(f"Date: {datetime.date.today().strftime('%B %d, %Y')}")
    doc.add_paragraph(f"Author: [Your Name]")
    doc.add_paragraph(f"OpenRouter Status: {'Operational' if check_openrouter_status() else 'Down'}")
    doc.add_paragraph(f"Mode: {'Deep Research' if deep_research else 'Quick Research'}")
    
    # Research Summary Section
    doc.add_heading("Research Summary", level=1)
    for item in data:
        p = doc.add_paragraph()
        p.add_run(f"• {item['title']}").bold = True
        doc.add_paragraph(item['content'])
        doc.add_paragraph(f"Source: {item['url']}")
        doc.add_paragraph()  # Add spacing

    # Process other sections
    sections = summary.split("\n\n")
    current_heading = None
    references = []
    
    for section in sections:
        if section.startswith("**") and section.endswith("**"):
            current_heading = section.strip("**").rstrip(":")
            if current_heading != "References":
                doc.add_heading(current_heading, level=2)
        else:
            if current_heading == "References":
                # Collect references for later processing
                refs = section.split("\n")
                for ref in refs:
                    if ref.strip():
                        references.append(ref.strip())
            elif current_heading == "Analysis":
                # Split into smaller paragraphs for readability
                sentences = re.split(r'(?<=[.!?])\s+', section)
                new_paragraphs = []
                current_paragraph = []
                sentence_count = 0
                
                max_sentences = 4 if len(sentences) < 15 else 6
                
                for sentence in sentences:
                    current_paragraph.append(sentence)
                    sentence_count += 1
                    if sentence_count >= max_sentences:
                        new_paragraphs.append(' '.join(current_paragraph))
                        current_paragraph = []
                        sentence_count = 0
                
                if current_paragraph:
                    new_paragraphs.append(' '.join(current_paragraph))
                
                for paragraph in new_paragraphs:
                    if paragraph.strip():
                        p = doc.add_paragraph(paragraph.strip())
                        p.paragraph_format.space_after = Pt(12)
            elif current_heading == "Key Findings":
                findings = re.split(r'(?=\d+\.)', section)
                for finding in findings:
                    if finding.strip():
                        p = doc.add_paragraph(finding.strip())
                        p.paragraph_format.space_after = Pt(6)
            else:
                formatted_text = re.sub(r"\*\*(.*?)\*\*", r"\1", section)
                p = doc.add_paragraph(formatted_text)
                p.paragraph_format.space_after = Pt(12)

    # Add References section at the end
    if references:
        doc.add_heading("References", level=2)
        
        for i, ref in enumerate(references, 1):
            formatted_ref = format_reference_for_pdf(ref)
            if formatted_ref:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.5)
                p.paragraph_format.space_after = Pt(12)
                p.add_run(f"{i}. {formatted_ref}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit app setup
st.title("AI agent-based Deep Research")
st.write("Enter a query to research and get a detailed response using Tavily and OpenRouter. Deep Research AI Agentic System that crawls websites using Tavily for online information gathering.")

# Sidebar: Status and Info
st.sidebar.header("OpenRouter Status")
if check_openrouter_status():
    st.sidebar.success("Operational", icon="✅")
else:
    st.sidebar.error("Down", icon="❌")

st.sidebar.header("About")
st.sidebar.write("Dual-AI-agent system using Tavily for research and OpenRouter for drafting with the model of your choosing.")
st.sidebar.write("Built with LangChain, LangGraph, and Streamlit Application.")

# User input with Deep Research toggle
query = st.text_input("Research Query", "Grammar Correction model using Machine Learning")
deep_research = st.checkbox("Deep Research Mode", value=False, help="Enable for a detailed, research-paper-style summary (5-6+ pages).")

# Research Settings Header
st.markdown("""
    <h2 style='color: #79c0ff; font-size: 1.8rem; font-weight: 600; margin-top: 1.8rem; margin-bottom: 1rem;'>
        🛠️ Research Settings
    </h2>
""", unsafe_allow_html=True)

# Create three columns for better organization
col1, col2, col3 = st.columns(3)

with col1:
    st.markdown("#### Writing Style")
    writing_style = st.radio(
        "Select writing style",
        options=["Academic", "Business", "Technical", "Casual"],
        index=0,
        key="writing_style_radio",
        help="Choose the tone and style of your research"
    )

with col2:
    st.markdown("#### Language")
    language = st.radio(
        "Select language",
        options=["English", "Spanish", "French", "German", "Chinese"],
        index=0,
        key="language_radio",
        help="Choose output language"
    )

with col3:
    st.markdown("#### Citation Format")
    citation_format = st.radio(
        "Select citation style",
        options=["APA", "MLA", "IEEE"],
        index=0,
        key="citation_format_radio",
        help="Choose citation formatting style"
    )

# Add word count slider below the columns
st.markdown("#### Target Word Count")
target_word_count = st.slider(
    "Select target word count",
    min_value=500,
    max_value=5000,
    value=1000,
    step=100,
    key="word_count_slider",
    help="Choose the approximate length of your research paper"
)

# Update all session state values
st.session_state.writing_style = writing_style
st.session_state.language = language
st.session_state.citation_format = citation_format
st.session_state.target_word_count = target_word_count

# Display current settings
st.markdown("""
    <div style='margin-top: 2rem;'>
        <p style='color: #58a6ff; font-size: 1.1rem; font-weight: 500;'>Current Settings:</p>
        <ul style='list-style-type: none; padding: 0; color: #c9d1d9;'>
            <li>📝 Writing Style: <strong>{}</strong></li>
            <li>📚 Citation Format: <strong>{}</strong></li>
            <li>🌐 Language: <strong>{}</strong></li>
            <li>📊 Target Words: <strong>{}</strong></li>
        </ul>
    </div>
""".format(
    writing_style.title(),
    citation_format,
    language.title(),
    target_word_count
), unsafe_allow_html=True)

# Style descriptions dictionary with capitalized keys
style_descriptions = {
    "Academic": "Formal scholarly writing with rigorous citations",
    "Business": "Professional tone with actionable insights",
    "Technical": "Detailed technical analysis and specifications",
    "Casual": "Accessible, conversational explanation"
}

with st.expander(" Writing Style Preview"):
    st.markdown(f"""
        <div style='background-color: #161b22; border: 1px solid #30363d; border-radius: 8px; padding: 1rem;'>
            <p style='color: #58a6ff; font-weight: bold;'>{writing_style}</p>
            <p style='color: #c9d1d9;'>{style_descriptions[writing_style]}</p>
        </div>
    """, unsafe_allow_html=True)
    st.markdown(f"<p style='color: #c9d1d9;'>Example citation in {citation_format}:</p>", unsafe_allow_html=True)
    example_citation = {
        "title": "Sample Research Paper",
        "url": "https://example.com/research",
    }
    st.code(format_citation(example_citation, citation_format))

# Research button logic
if st.button("Run Research"):
    if not query.strip():
        st.error("Please enter your research query.")
    elif not check_openrouter_status():
        st.error("OpenRouter is currently down. Please try again later.")
    else:
        try:
            with st.spinner("Processing your request..."):
                # Initialize progress bar and status
                progress_bar = st.progress(0)
                status_text = st.empty()

                # Step 1: Fetch research data
                status_text.text("Step 1/3: Fetching research data... 🔍")
                logging.info(f"Starting research for query: {query}, deep_research: {deep_research}, target_word_count: {target_word_count}")
                research_data, response = run_research(
                    query,
                    deep_research=deep_research,
                    target_word_count=target_word_count,
                    writing_style=writing_style,
                    citation_format=citation_format,
                    language=language
                )
                progress_bar.progress(33)

                # Step 2: Drafting response
                status_text.text("Step 2/3: Drafting response... ")
                if "Error drafting response" in response:
                    st.error(response)
                    logging.error(f"Failed to draft response: {response}")
                else:
                    progress_bar.progress(66)

                    # Step 3: Generating PDF
                    status_text.text("Step 3/3: Generating PDF report... ")
                    st.success("Research completed! 🎉", icon="✅")
                    st.subheader("Structured Summary 📝")

                    # Display summary with collapsible sections
                    sections = response.split("\n\n")
                    current_section = None
                    section_content = []

                    # First, display Research Summary section
                    st.write("### Research Summary 📊")
                    with st.expander("View Research Summary", expanded=True):
                        for item in research_data:
                            st.markdown(f"**{item['title']}**")
                            st.write(item['content'])
                            st.markdown(f"[Source]({item['url']})")
                            st.markdown("---")

                    st.write("### Detailed Analysis 📝")
                    
                    # Then process other sections
                    for section in sections:
                        if section.startswith("**") and section.endswith("**"):
                            # If we have accumulated content for the previous section, display it
                            if current_section and section_content:
                                with st.expander(current_section, expanded=True):
                                    if current_section == "References":
                                        # Special handling for references with numbering
                                        references = "\n".join(section_content).strip().split("\n")
                                        for i, ref in enumerate(references, 1):
                                            if ref.strip():
                                                # Format reference with number and add horizontal line
                                                st.markdown(f"{i}. {ref.strip()}")
                                                st.markdown("---")
                                    elif current_section == "Analysis":
                                        # Keep existing Analysis formatting
                                        analysis_text = "\n".join(section_content)
                                        sentences = re.split(r'(?<=[.!?])\s+', analysis_text)
                                        paragraphs = []
                                        current_paragraph = []
                                        sentence_count = 0
                                        
                                        for sentence in sentences:
                                            current_paragraph.append(sentence)
                                            sentence_count += 1
                                            if sentence_count >= 6:
                                                paragraphs.append(" ".join(current_paragraph))
                                                current_paragraph = []
                                                sentence_count = 0
                                        
                                        if current_paragraph:
                                            paragraphs.append(" ".join(current_paragraph))
                                        
                                        for paragraph in paragraphs:
                                            st.write(paragraph)
                                            st.write("")
                                    else:
                                        for content in section_content:
                                            st.markdown(content)
                                            
                            # Start a new section
                            current_section = section.strip("**").rstrip(":")
                            section_content = []
                        else:
                            if current_section:
                                section_content.append(section)

                    # Display the last section
                    if current_section and section_content:
                        with st.expander(current_section, expanded=True):
                            if current_section == "References":
                                references = "\n".join(section_content).strip().split("\n")
                                for i, ref in enumerate(references, 1):
                                    if ref.strip():
                                        # Format reference with number and add horizontal line
                                        st.markdown(f"{i}. {ref.strip()}")
                                        st.markdown("---")
                            elif current_section == "Analysis":
                                # Keep existing Analysis formatting
                                analysis_text = "\n".join(section_content)
                                sentences = re.split(r'(?<=[.!?])\s+', analysis_text)
                                paragraphs = []
                                current_paragraph = []
                                sentence_count = 0
                                
                                for sentence in sentences:
                                    current_paragraph.append(sentence)
                                    sentence_count += 1
                                    if sentence_count >= 6:
                                        paragraphs.append(" ".join(current_paragraph))
                                        current_paragraph = []
                                        sentence_count = 0
                                
                                if current_paragraph:
                                    paragraphs.append(" ".join(current_paragraph))
                                
                                for paragraph in paragraphs:
                                    st.write(paragraph)
                                    st.write("")
                            else:
                                for content in section_content:
                                    st.markdown(content)

                    # Calculate word count and page 
                    word_count = len(response.split())
                    page_estimate = word_count // 400 + 1  # Rough estimate: ~400 words per page
                    st.info(f"Summary contains {word_count} words, estimated at {page_estimate} pages.")

                    # Store results in session state
                    st.session_state.research_data = research_data
                    st.session_state.response = response
                    st.session_state.pdf_buffer = generate_pdf(query, research_data, response, deep_research=deep_research)
                    st.session_state.word_buffer = generate_docx(query, research_data, response, deep_research=deep_research)

                    #  Display Interactive Research Data
                    st.write("### Research Data 📚")
                    for item in research_data:
                        with st.expander(item['title']):
                            st.write(item['content'])
                            st.markdown(f"[Visit Source]({item['url']})")

        except Exception as e:
            st.error(f"Failed after retries: {str(e)}")
            logging.error(f"Failed to process query '{query}': {str(e)}")
        finally:
            progress_bar.empty()  # Clear progress bar
            status_text.empty()  # Clear status text

# Display download options if research data is available
if st.session_state.research_data and st.session_state.response:
    st.write("### Download Options")
    
    # Simple format selection without session state
    selected_format = st.radio(
        "Select format:",
        ["PDF (Recommended)", "Word", "Markdown", "Text"],
        horizontal=True
    )

    # Show description based on file selection
    if selected_format == "PDF (Recommended)":
        st.caption("Download as a PDF file.")
        st.download_button(
            label="Download PDF 📥",
            data=st.session_state.pdf_buffer,
            file_name="research_report.pdf",
            mime="application/pdf"
        )
    elif selected_format == "Word":
        st.caption("Download as a Word document.")
        st.download_button(
            label="Download Word 📥",
            data=st.session_state.word_buffer,
            file_name="research_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    elif selected_format == "Markdown":
        st.caption("Download as a Markdown file.")
        st.download_button(
            label="Download Markdown 📥",
            data=st.session_state.response,
            file_name="research_summary.md",
            mime="text/markdown"
        )
    else:  # Text
        st.caption("Download as plain text.")
        st.download_button(
            label="Download Text 📥",
            data=st.session_state.response,
            file_name="research_summary.txt",
            mime="text/plain"
        )

# Feedback form in sidebar
st.sidebar.header("Feedback")
feedback = st.sidebar.text_area("Please give us feedback on how can we improve? (Optional)")
if st.sidebar.button("Submit Feedback"):
    st.sidebar.success("Thank you for your feedback! 🙏")
    with open("feedback.txt", "a") as f:
        f.write(f"{feedback}\n")